"""
docx_utils.py
=============
Generates a professionally formatted Microsoft Word (.docx) examination
document with embedded Matplotlib diagrams, marks distribution table,
styled question layout, and answer-space ruling lines.
"""

import os
import re
from typing import List, Dict, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ─────────────────────────── COLOUR PALETTE ───────────────────────────────
C_DARK   = RGBColor(15,  23,  42)    # #0f172a — near-black
C_BLUE   = RGBColor(30,  64, 175)    # #1e40af — question numbers
C_ACCENT = RGBColor( 2, 132, 199)    # #0284c7 — marks badge
C_MUTED  = RGBColor(71,  85, 105)    # #475569 — meta / subtitles
C_BODY   = RGBColor(30,  41,  59)    # #1e293b — body text
C_RULE   = RGBColor(203, 213, 225)   # #cbd5e1 — divider lines
C_HDR_BG = RGBColor(15,  23,  42)    # table header background
C_WHITE  = RGBColor(255, 255, 255)


# ─────────────────────────── Utility helpers ──────────────────────────────
def _set_cell_bg(cell, hex_colour: str):
    """Fills a table cell background with a hex colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour.lstrip("#"))
    tcPr.append(shd)


def _set_cell_border(cell, *, top=None, bottom=None, left=None, right=None):
    """Adds borders to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, colour in (("top", top), ("bottom", bottom),
                         ("left", left), ("right", right)):
        if colour:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), colour.lstrip("#"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_horizontal_rule(doc: Document, colour_hex: str = "#cbd5e1"):
    """Inserts a styled horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr  = OxmlElement("w:bottom")
    bdr.set(qn("w:val"),   "single")
    bdr.set(qn("w:sz"),    "6")
    bdr.set(qn("w:space"), "1")
    bdr.set(qn("w:color"), colour_hex.lstrip("#"))
    pBdr.append(bdr)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def _add_answer_lines(doc: Document, marks: int):
    """Inserts ruled dotted lines proportional to question marks."""
    n_lines = min(max(3, marks // 3), 10)
    for _ in range(n_lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run("_" * 105)
        run.font.color.rgb = C_RULE
        run.font.size = Pt(8)


# ─────────────────────────── LaTeX cleaner ────────────────────────────────
def _clean_latex_for_docx(text: str) -> str:
    """Converts LaTeX math markup to clean readable Unicode for Word."""
    if not text:
        return ""

    # Remove figure environments
    text = re.sub(r"\\begin\{figure\}[\s\S]*?\\end\{figure\}", "", text)
    text = re.sub(r"\\includegraphics(\[.*?\])?\{.*?\}", "", text)
    text = re.sub(r"\\caption\{.*?\}", "", text)
    text = re.sub(r"\\centering", "", text)
    text = re.sub(r"\\label\{.*?\}", "", text)

    replacements = [
        (r"\\frac\{([^}]+)\}\{([^}]+)\}", r"(\1 / \2)"),
        (r"\\sqrt\{([^}]+)\}", r"√(\1)"),
        (r"\\sqrt", "√"),
        (r"\\partial", "∂"),   (r"\\nabla",   "∇"),
        (r"\\int_\{?([^}^ ]+)\}?\^\{?([^}^ ]+)\}?", r"∫[\1 to \2]"),
        (r"\\int",    "∫"),    (r"\\oint",    "∮"),
        (r"\\sum_\{?([^}^ ]+)\}?\^\{?([^}^ ]+)\}?", r"∑[\1 to \2]"),
        (r"\\sum",    "∑"),    (r"\\prod",    "∏"),
        (r"\\alpha",  "α"),    (r"\\beta",    "β"),
        (r"\\gamma",  "γ"),    (r"\\delta",   "δ"),
        (r"\\epsilon","ε"),    (r"\\zeta",    "ζ"),
        (r"\\eta",    "η"),    (r"\\theta",   "θ"),
        (r"\\lambda", "λ"),    (r"\\mu",      "μ"),
        (r"\\nu",     "ν"),    (r"\\xi",      "ξ"),
        (r"\\pi",     "π"),    (r"\\rho",     "ρ"),
        (r"\\sigma",  "σ"),    (r"\\tau",     "τ"),
        (r"\\phi",    "φ"),    (r"\\chi",     "χ"),
        (r"\\psi",    "ψ"),    (r"\\omega",   "ω"),
        (r"\\Gamma",  "Γ"),    (r"\\Delta",   "Δ"),
        (r"\\Theta",  "Θ"),    (r"\\Lambda",  "Λ"),
        (r"\\Pi",     "Π"),    (r"\\Sigma",   "Σ"),
        (r"\\Phi",    "Φ"),    (r"\\Psi",     "Ψ"),
        (r"\\Omega",  "Ω"),
        (r"\\infty",  "∞"),    (r"\\pm",      "±"),
        (r"\\times",  "×"),    (r"\\div",     "÷"),
        (r"\\cdot",   "·"),    (r"\\leq",     "≤"),
        (r"\\geq",    "≥"),    (r"\\neq",     "≠"),
        (r"\\approx", "≈"),    (r"\\equiv",   "≡"),
        (r"\\propto", "∝"),    (r"\\to",      "→"),
        (r"\\rightarrow","→"), (r"\\leftarrow","←"),
        (r"\\Rightarrow","⇒"), (r"\\forall",  "∀"),
        (r"\\exists", "∃"),    (r"\\in",      "∈"),
        (r"\\notin",  "∉"),    (r"\\subset",  "⊂"),
        (r"\\cup",    "∪"),    (r"\\cap",     "∩"),
        (r"\\perp",   "⊥"),    (r"\\parallel","∥"),
        (r"\\angle",  "∠"),    (r"\\triangle","△"),
        (r"\\sin",    "sin"),  (r"\\cos",     "cos"),
        (r"\\tan",    "tan"),  (r"\\exp",     "exp"),
        (r"\\log",    "log"),  (r"\\ln",      "ln"),
        (r"\\mathbf\{([^}]+)\}", r"\1"),
        (r"\\textbf\{([^}]+)\}", r"\1"),
        (r"\\textit\{([^}]+)\}", r"\1"),
        (r"\\emph\{([^}]+)\}",   r"\1"),
        (r"\\mathbb\{R\}\^?(\d+)?", r"ℝ\1"),
        (r"\\mathbb\{([^}]+)\}", r"\1"),
        (r"\\mathrm\{([^}]+)\}", r"\1"),
        (r"\\text\{([^}]+)\}",   r"\1"),
        (r"\^\{([^}]+)\}",  r"^\1"),
        (r"_\{([^}]+)\}",   r"_\1"),
    ]

    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text)

    # Strip math delimiters
    text = re.sub(r"\\\((.*?)\\\)", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\\\[(.*?)\\\]", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\$\$([^\$]+)\$\$", r"\1", text)
    text = re.sub(r"\$([^\$]+)\$",     r"\1", text)

    # Strip remaining LaTeX commands
    text = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+", "", text)

    return text.strip()


# ─────────────────────────── Main DOCX builder ────────────────────────────
def create_docx(
    exam_title: str,
    questions: List[Dict[str, Any]],
    max_marks: int,
    output_path: str,
    institution: str = "University Examination Board",
    course_code: str = "",
    exam_year: str = "2026",
) -> str:
    """
    Generates a styled Microsoft Word exam document with:
    • Institution header + title + metadata
    • General instructions paragraph
    • Marks distribution table
    • Per-question blocks: number, marks, text, diagram, answer lines
    • Consistent Calibri typography throughout
    """
    doc      = Document()
    base_dir = os.path.dirname(os.path.abspath(__file__))

    # ── Document margins ───────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ── Institution name ───────────────────────────────────────────────────
    inst_p = doc.add_paragraph()
    inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_p.paragraph_format.space_after = Pt(2)
    ir = inst_p.add_run(institution.upper())
    ir.font.name  = "Calibri"
    ir.font.size  = Pt(9)
    ir.font.color.rgb = C_MUTED
    ir.font.bold  = False

    # ── Exam Title ─────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    tr = title_p.add_run(exam_title.upper())
    tr.bold           = True
    tr.font.size      = Pt(18)
    tr.font.name      = "Calibri"
    tr.font.color.rgb = C_DARK

    # ── Meta line ──────────────────────────────────────────────────────────
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(8)
    course_str = f"Course: {course_code}   |   " if course_code else ""
    mr = meta_p.add_run(
        f"{course_str}"
        f"Year: {exam_year}   |   "
        f"Maximum Marks: {max_marks}   |   "
        f"Total Questions: {len(questions)}   |   "
        f"Time: 3 Hours"
    )
    mr.font.name      = "Calibri"
    mr.font.size      = Pt(9.5)
    mr.font.color.rgb = C_MUTED

    # ── Bold divider ───────────────────────────────────────────────────────
    _add_horizontal_rule(doc, "#0f172a")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── General Instructions ───────────────────────────────────────────────
    inst_block = doc.add_paragraph()
    inst_block.paragraph_format.space_after = Pt(10)
    ib = inst_block.add_run("GENERAL INSTRUCTIONS:  ")
    ib.bold           = True
    ib.font.size      = Pt(10)
    ib.font.name      = "Calibri"
    ib.font.color.rgb = C_DARK
    ic = inst_block.add_run(
        "Answer all questions. Show all working, derivations and calculations. "
        "Draw and label diagrams wherever required. Use of scientific calculator is permitted."
    )
    ic.italic         = True
    ic.font.size      = Pt(9.5)
    ic.font.name      = "Calibri"
    ic.font.color.rgb = C_MUTED

    # ── Marks Distribution Table ───────────────────────────────────────────
    n_q      = len(questions)
    col_count = n_q + 2      # Q.No | Q1 … Qn | Total
    mdist    = doc.add_table(rows=2, cols=col_count)
    mdist.alignment = WD_TABLE_ALIGNMENT.CENTER
    mdist.style = "Table Grid"

    # Header row
    headers = ["Q.No"] + [str(i) for i in range(1, n_q + 1)] + ["Total"]
    for ci, hdr in enumerate(headers):
        cell = mdist.cell(0, ci)
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold      = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.name = "Calibri"
        cell.paragraphs[0].runs[0].font.color.rgb = C_WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "#0f172a")

    # Marks row
    mark_vals = ["Marks"]
    for q in questions:
        q_d = dict(q) if isinstance(q, dict) else q.__dict__
        mark_vals.append(str(q_d.get("marks", 0)))
    mark_vals.append(str(max_marks))

    for ci, val in enumerate(mark_vals):
        cell = mdist.cell(1, ci)
        cell.text = val
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = C_DARK
        if ci == 0:
            run.bold = True
        if ci == col_count - 1:
            run.bold = True
            run.font.color.rgb = C_BLUE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        bg = "#f1f5f9" if ci % 2 == 0 else "#ffffff"
        if ci == col_count - 1:
            bg = "#dbeafe"
        _set_cell_bg(cell, bg)

    # Set uniform column widths
    tbl_width_cm = 17.0
    col_w_cm = tbl_width_cm / col_count
    for row in mdist.rows:
        for cell in row.cells:
            cell.width = Cm(col_w_cm)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Questions ──────────────────────────────────────────────────────────
    fig_counter = 0
    for idx, q in enumerate(questions, start=1):
        q_dict   = dict(q) if isinstance(q, dict) else q.__dict__
        marks    = int(q_dict.get("marks", 0))
        raw_text = q_dict.get("text", "")
        clean    = _clean_latex_for_docx(raw_text)
        img_path = q_dict.get("image_path")

        # ── Question header row ────────────────────────────────────────────
        # Use a 2-column table for left-justified Q# and right-justified marks
        hdr_tbl = doc.add_table(rows=1, cols=2)
        hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr_tbl.style     = "Table Grid"
        hdr_tbl.rows[0].cells[0].merge(hdr_tbl.rows[0].cells[0])

        # Left cell: "Question X"
        lc = hdr_tbl.cell(0, 0)
        lc.text = ""
        lp  = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after  = Pt(0)
        lr  = lp.add_run(f"Question {idx}")
        lr.bold           = True
        lr.font.size      = Pt(11)
        lr.font.name      = "Calibri"
        lr.font.color.rgb = C_BLUE
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_bg(lc, "#eff6ff")

        # Right cell: marks badge
        rc = hdr_tbl.cell(0, 1)
        rc.text = ""
        rp  = rc.paragraphs[0]
        rp.paragraph_format.space_before = Pt(0)
        rp.paragraph_format.space_after  = Pt(0)
        rr  = rp.add_run(f"[{marks} Marks]")
        rr.bold           = True
        rr.font.size      = Pt(10)
        rr.font.name      = "Calibri"
        rr.font.color.rgb = C_ACCENT
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_cell_bg(rc, "#eff6ff")

        # Column widths (approx A4 page with 0.9" margins → ~16.8 cm usable)
        hdr_tbl.cell(0, 0).width = Cm(12.5)
        hdr_tbl.cell(0, 1).width = Cm(4.3)

        # Remove table borders (clean look — blue bottom border only)
        for cell in (lc, rc):
            _set_cell_border(cell, bottom="#bfdbfe")

        # ── Question body ──────────────────────────────────────────────────
        body_p = doc.add_paragraph()
        body_p.paragraph_format.space_before = Pt(5)
        body_p.paragraph_format.space_after  = Pt(4)
        body_p.paragraph_format.left_indent  = Inches(0.15)
        br = body_p.add_run(clean)
        br.font.size      = Pt(11)
        br.font.name      = "Calibri"
        br.font.color.rgb = C_BODY

        # ── Diagram ────────────────────────────────────────────────────────
        if img_path:
            abs_img = os.path.abspath(os.path.join(base_dir, img_path))
            if os.path.exists(abs_img):
                fig_counter += 1
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(4)
                img_p.paragraph_format.space_after  = Pt(2)
                try:
                    img_p.add_run().add_picture(abs_img, width=Inches(4.6))
                except Exception as ie:
                    print(f"[DOCX Export] Warning adding picture: {ie}")

                cap_p = doc.add_paragraph(f"Figure {fig_counter}")
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.paragraph_format.space_after = Pt(4)
                cap_r = cap_p.runs[0]
                cap_r.italic      = True
                cap_r.font.size   = Pt(9)
                cap_r.font.name   = "Calibri"
                cap_r.font.color.rgb = C_MUTED

        # ── Answer lines ───────────────────────────────────────────────────
        _add_answer_lines(doc, marks)

        # ── Divider ────────────────────────────────────────────────────────
        _add_horizontal_rule(doc, "#e2e8f0")
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)

        # Page break every 5 questions to avoid overcrowding
        if idx % 5 == 0 and idx < len(questions):
            doc.add_page_break()

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path
